from flask import Flask, render_template, request, send_file,jsonify, session, redirect, url_for,send_from_directory
from moviepy.editor import VideoFileClip, AudioFileClip, CompositeVideoClip
from moviepy.editor import VideoFileClip, concatenate_audioclips
from moviepy.editor import VideoFileClip
from moviepy.audio.io.AudioFileClip import AudioFileClip
from record_camera import VideoRecorder
import a
import movie
import threading
from PIL import Image
import os
import fitz  # PyMuPDF for PDF handling
from PIL import Image
from docx import Document
import os
from pyppeteer import launch
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from werkzeug.utils import secure_filename
import datetime

app = Flask(__name__)
app.secret_key = 'meni_the_coder'

VIDEO_PATH = 'movie.avi'
AUDIO_PATH = 'audio.wav'
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
def create_audio(seconds):
    a.create_audio(seconds)

def create_movie(seconds):
    movie.create_movie(seconds)

def resize_images_to_equal_size(image1_path, image2_path):
    img1 = Image.open(image1_path)
    img2 = Image.open(image2_path)

    # Get the sizes of both images
    width1, height1 = img1.size
    width2, height2 = img2.size

    # Calculate the minimum dimensions to resize both images
    min_width = max(width1, width2)
    min_height = max(height1, height2)

    # Resize both images to have equal dimensions
    img1_resized = img1.resize((min_width, min_height))
    img2_resized = img2.resize((min_width, min_height))

    return img1_resized, img2_resized
def image1_dominate(img1, width1, height1, width_ratio, height_ratio):
    new_width1 = int(width1 * width_ratio)
    new_height1 = int(height1 * height_ratio)
    img1 = img1.resize((new_width1, new_height1))
    return img1

def image2_dominate(img2, width2, height2, width_ratio, height_ratio):
    new_width2 = int(width2 * width_ratio)
    new_height2 = int(height2 * height_ratio)
    img2 = img2.resize((new_width2, new_height2))
    return img2

#############################
#function for covert
#############################
def pdf_to_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    images = []

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        image = page.get_pixmap()
        img = Image.frombytes("RGB", [image.width, image.height], image.samples)
        images.append(img)

    pdf_document.close()
    return images

def save_images(images, output_format='PNG'):
    download_path = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Ensure the directory exists, create it if necessary
    if not os.path.exists(download_path):
        os.makedirs(download_path)

    for i, img in enumerate(images):
        img.save(os.path.join(download_path, f"output_{i}.{output_format.lower()}"), format=output_format)

def pdf_to_text(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ''
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    pdf_document.close()
    return text

def text_to_word(text, output_path='output.docx'):
    download_path = os.path.join(os.path.expanduser("~"), "Downloads", output_path)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(download_path)

###################################
# word to other
####################################
def convert_docx_to_pdf(input_docx, output_pdf):
    doc = Document(input_docx)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
    </head>
    <body>
        {text}
    </body>
    </html>
    '''

    async def generate_pdf():
        browser = await launch()
        page = await browser.newPage()
        await page.setContent(html)
        await page.pdf({'path': output_pdf, 'format': 'A4'})
        await browser.close()

    import asyncio
    asyncio.get_event_loop().run_until_complete(generate_pdf())

def convert_docx_to_jpeg(input_docx, output_, format_type):
    # Read the DOCX file
    doc = Document(input_docx)

    # Create an image
    img_width = 800
    img_height = 600
    background_color = 'white'
    font_size = 20
    font_color = 'black'

    image = Image.new('RGB', (img_width, img_height), color=background_color)
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", font_size)

    # Convert the text content from the DOCX to image
    text_content = '\n'.join([para.text for para in doc.paragraphs])
    draw.text((10, 10), text_content, font=font, fill=font_color)
    output_t = os.path.join(os.path.expanduser("~"), "Downloads", output_)
    # Save the image as JPEG
    image.save(output_t, format_type)

def convert_png_to_pdf(input_path, output_filename):
    image = Image.open(input_path)

    # Get the path to the user's download directory
    download_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Create the full output path including the download directory
    output_path = os.path.join(download_dir, output_filename)

    # Create a PDF
    pdf = canvas.Canvas(output_path, pagesize=image.size)

    # Draw the PNG image onto the PDF
    pdf.drawImage(input_path, 0, 0, width=image.width, height=image.height)

    # Save the PDF file
    pdf.save()
    print(f"Converted {input_path} to {output_path}")

def convert_png_to_jpeg(input_path, output_path):
    try:
        # Open the PNG image
        with Image.open(input_path) as img:
            # Convert PNG to JPEG
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Save the image as JPEG
            img.save(output_path, 'JPEG')

            print(f"File saved successfully as {output_path}")
    except Exception as e:
        print(f"Error converting file: {e}")

def convert_jpeg_to_png(input_path, output_path):
    try:
        # Open the JPEG image
        with Image.open(input_path) as img:
            # Convert JPEG to PNG
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Save the image as PNG
            img.save(output_path, 'PNG')

            print(f"File saved successfully as {output_path}")
    except Exception as e:
        print(f"Error converting file: {e}")

def convert_jpeg_to_pdf(input_image, output_pdf):
    # Create a new PDF file
    c = canvas.Canvas(output_pdf, pagesize=letter)

    # Open the image using PIL
    img = Image.open(input_image)

    # Get the dimensions of the image
    width, height = img.size

    # Add the image to the PDF file
    c.setPageSize((width, height))
    c.drawImage(input_image, 0, 0, width, height)
    c.showPage()

    # Save the PDF file
    c.save()
@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'), 'favicon.ico', mimetype='image/png')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/thankyou')
def thanks():
    return render_template('thank_you.html')

@app.route('/process', methods=['POST'])
def process():
    try:
        user_input = float(request.form['minutes'])
        time = user_input * 60

        # Rest of your code remains the same
        # ...
    except KeyError:
        return "Invalid input data. Please enter a valid number of minutes."
    audio_thread = threading.Thread(target=create_audio, args=(time,))
    movie_thread = threading.Thread(target=create_movie, args=(time,))



    audio_thread.start()
    movie_thread.start()


    audio_thread.join()
    movie_thread.join()



    try:
        video = VideoFileClip('movie.avi')

        # Replace 'audio.wav' with the path to your WAV audio file
        wav = AudioFileClip('audio.wav')

        # Set the audio of the video file to the provided audio
        video = video.set_audio(wav)

        downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

        # Define the output file path within the downloads directory
        output_file_path = os.path.join(downloads_dir, 'output_video.mp4')

        # Export the final video with the combined audio as an MP4 file in the downloads folder
        video.write_videofile(output_file_path, codec='libx264', audio_codec='aac')

        return redirect(url_for('thanks'))
    except Exception as e:
        return f'Error during video processing: {str(e)}'


@app.route('/process_with_camera', methods=['POST'])
def process_with_camera():
    try:
        user_input = float(request.form['minutes'])
        time = user_input * 60

        # Rest of your code remains the same
        # ...
    except KeyError:
        return "Invalid input data. Please enter a valid number of minutes."
    audio_thread = threading.Thread(target=create_audio, args=(time,))
    movie_thread = threading.Thread(target=create_movie, args=(time,))


    video_recorder = VideoRecorder(time)
    video_recorder.start()

    while video_recorder.get_time_diff() is None:
        # Perform actions while wait_time is None
        pass

    print("start")
    audio_thread.start()

    movie_thread.start()


    audio_thread.join()

    movie_thread.join()

    video_recorder.join()
    print("all files done")


    try:

        print("1")
        video = VideoFileClip('movie.avi')
        print("2")
        # Replace 'audio.wav' with the path to your WAV audio file
        wav = AudioFileClip('audio.wav')
        print("3")

        # Set the audio of the video file to the provided audio
        video = video.set_audio(wav)
        print("4")

        camera_feed = VideoFileClip('camera_video.mp4')
        print("5")

        # Define the size of the final video
        final_width = video.w + camera_feed.w
        final_height = max(video.h, camera_feed.h)

        # Resize camera feed to fit the bottom left corner of the video
        camera_feed = camera_feed.resize(height=video.h / 4)  # Adjust the size as needed
        print("5")
        # Place the camera feed at the bottom left corner
        camera_position = (0, final_height - camera_feed.h)
        print("6")
        # Combine the existing video and the camera feed as an overlay
        final_video = CompositeVideoClip([video.set_position('center'), camera_feed.set_position(camera_position)])

        # Define the output file path within the downloads directory
        downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')
        output_file_path = os.path.join(downloads_dir, 'output_video_with_camera.mp4')

        # Export the final video with the combined audio and camera overlay as an MP4 file in the downloads folder
        final_video.write_videofile(output_file_path, codec='libx264', audio_codec='aac')

        return redirect(url_for('thanks'))
    except Exception as e:
        return f'Error during video processing: {str(e)}'
@app.route('/coordinate', methods=['POST'])
def handle_coordinate():
    if request.method == 'POST' and request.is_json:
        data = request.json  # Extract JSON data from the request
        x = data.get('x')
        y = data.get('y')
        print(x,y)
        session['x'] = x
        session['y'] = y

        # Process the received 'x' and 'y' coordinates (Perform actions as needed)
        # For example, you can perform calculations, database operations, etc.

        # Return a response (This is just an example)
        response_data = {'message': 'Received coordinates successfully', 'x': x, 'y': y}
        return jsonify(response_data), 200
    else:
        return jsonify({'error': 'Invalid request or Content-Type is not application/json'}), 400

@app.route('/merge_images', methods=['POST'])
def merge_images_custom():
    image1_path = request.files['file1']
    image2_path = request.files['file2']
    output_path = request.form['image_format']

    width_ratio = float(session['x'])
    height_ratio = float(session['y'])
    if width_ratio>0.9:
        width_ratio=1

    if  height_ratio>0.9:
        height_ratio=1


    if width_ratio != 1 and height_ratio != 1:
        dominate = int(request.form['dominate'])  # If 'dominate' not provided, default to 1
        if dominate==1:
            dominate=2
        elif dominate==2:
            dominate=1
    else:
        dominate = 1

    corner = int(request.form['corner'])

    output_filename = (f"merged_image.{output_path}")
    download_path = os.path.join(os.path.expanduser("~"), "Downloads", output_filename)
    img1_resized, img2_resized = resize_images_to_equal_size(image1_path, image2_path)

    width1, height1 = img1_resized.size
    width2, height2 = img2_resized.size
    wi = 0
    hi = 0

    if dominate == 1:
        if width_ratio == 1 and height_ratio != 1:
            print((1 - height_ratio))
            new_height1 = int(height1 * (1 - height_ratio))
            img1_resized = img1_resized.resize((width1, new_height1))
            hi = int(height2 * height_ratio)

        if width_ratio != 1 and height_ratio == 1:
            new_width1 = int(width1 * (1 - width_ratio))
            img1_resized = img1_resized.resize((new_width1, height1))
            wi = int(width2 * width_ratio)

        img2 = image1_dominate(img2_resized, width2, height2, width_ratio, height_ratio)

        merged_image = Image.new('RGB', (width1, height1))

        if corner == 1:
            print(corner)
            merged_image.paste(img1_resized, (wi, hi))
            merged_image.paste(img2, (0, 0))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 2:
            merged_image.paste(img1_resized, (0, hi))
            position_x = int(width1 * (1 - width_ratio))
            merged_image.paste(img2, (position_x, 0))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 3:
            merged_image.paste(img1_resized, (0, 0))
            position_x = int(width1 * (1 - width_ratio))
            position_y = int(height1 * (1 - height_ratio))
            merged_image.paste(img2, (position_x, position_y))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 4:
            merged_image.paste(img1_resized, (wi, 0))
            position_y = int(height1 * (1 - height_ratio))
            merged_image.paste(img2, (0, position_y))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")






    elif dominate == 2:
        img1 = image2_dominate(img1_resized, width1, height1, width_ratio, height_ratio)
        merged_image = Image.new('RGB', (width2, height2))
        merged_image.paste(img2_resized, (0, 0))
        if corner == 1:
            print(corner)
            merged_image.paste(img1, (0, 0))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 2:
            position_x = int(width1 * (1 - width_ratio))
            merged_image.paste(img1, (position_x, 0))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 3:
            position_x = int(width1 * (1 - width_ratio))
            position_y = int(height1 * (1 - height_ratio))
            merged_image.paste(img1, (position_x, position_y))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
        elif corner == 4:
            position_y = int(height1 * (1 - height_ratio))
            merged_image.paste(img1, (0, position_y))
            merged_image.save(download_path)
            print(f"Images merged with custom proportions and saved to {output_path}")
    return redirect(url_for('thanks'))
@app.route('/download_output')
def download_output():
    return send_file('output_video.mp4', as_attachment=True)


@app.route('/handle_image_info', methods=['POST'])
def handle_image_info():
    # Get the JSON data from the request
    data = request.get_json()

    # Check if valid data and imageId are present
    if data and 'imageId' in data:
        image_id = data['imageId']

        # Store the imageId in the session for further processing or tracking
        session['imageId'] = image_id

        # Perform actions specific to handling the imageId
        # Your logic for handling imageId here...

        # Return a success message along with the stored imageId
        return jsonify({'message': f'Information received for Image {image_id}', 'imageId': image_id}), 200
    else:
        # If no valid imageId is received, return an error
        return jsonify({'error': 'No valid imageId received'}), 400
@app.route('/convert', methods=['GET', 'POST'])
def upload_file():

    if request.method == 'POST':
        # Get the file and format from the form

        file = request.files['file']

        if file.filename != '':
            filename = secure_filename(file.filename)
            pdf_file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(pdf_file_path)


            format = session['imageId']
            print(format)


            if pdf_file_path.lower().endswith('.pdf'):


                if format == "PNG" or format == "JPEG":

                    # Convert PDF to images (PNG by default)
                    pdf_images = pdf_to_images(pdf_file_path)
                    if format == "PNG":
                        # Save images in different formats (PNG, JPEG)
                        save_images(pdf_images, output_format='PNG')
                    elif format == "JPEG":
                        save_images(pdf_images, output_format='JPEG')
                elif format == "DOCX":
                    extracted_text = pdf_to_text(pdf_file_path)
                    text_to_word(extracted_text)


            ###################################
            # word to other
            ####################################

            elif pdf_file_path.lower().endswith('.docx'):
                output_image_file = 'convert{}.png'

                if format == "PDF":
                    output_pdf_file = 'converted.pdf'  # Replace with your desired output PDF filename
                    download_path = os.path.join(os.path.expanduser("~"), "Downloads", output_pdf_file)
                    convert_docx_to_pdf(pdf_file_path, download_path)

                elif format == "PNG":
                    output_png_path = 'output_image.png'
                    convert_docx_to_jpeg(pdf_file_path, output_png_path, 'png')

                elif format == "JPEG":
                    output_jpeg_path = 'output_image.jpg'
                    convert_docx_to_jpeg(pdf_file_path, output_jpeg_path, 'jpeg')
            ###################################
            # png to other
            ####################################
            elif pdf_file_path.lower().endswith('.png'):
                if format == "PDF":
                    output_file = 'output.pdf'
                    convert_png_to_pdf(pdf_file_path, output_file)
                elif format == "JPEG":
                    download_directory = os.path.expanduser('~') + '/Downloads/'
                    output_file_path = os.path.join(download_directory, 'output_file.jpg')

                    # Convert PNG to JPEG and save in the download directory
                    convert_png_to_jpeg(pdf_file_path, output_file_path)

            elif pdf_file_path.lower().endswith(('.jpg', '.jpeg')):
                if format == "PDF":
                    # Get the path to the Downloads directory
                    download_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

                    # Replace 'output_pdf_file.pdf' with the desired output PDF file name
                    output_pdf_file = os.path.join(download_dir, 'output_pdf_file.pdf')

                    # Convert the file extension to lowercase for comparison

                    convert_jpeg_to_pdf(pdf_file_path, output_pdf_file)
                if format == "PNG":
                    downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')

                    # Input JPEG file path

                    # Output PNG file path in the Downloads directory
                    output_png_path = os.path.join(downloads_dir, 'converted_image.png')

                    # Call the function to convert JPEG to PNG
                    convert_jpeg_to_png(pdf_file_path, output_png_path)
    if os.path.exists(pdf_file_path):
        os.remove(pdf_file_path)
    return redirect(url_for('thanks'))

if __name__ == '__main__':
    app.run(debug=True)
